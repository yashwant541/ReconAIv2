from __future__ import annotations
import csv, hashlib, io
from dataclasses import dataclass
from typing import Protocol
import pandas as pd
from docx import Document as WordDocument
import pdfplumber
from .domain import FinancialCell, FinancialDocument, FinancialRow, FinancialSection, FinancialTable, Provenance
from .normalization import text

class Parser(Protocol):
    def parse(self, content: bytes, filename: str) -> FinancialDocument: ...

def _id(content: bytes) -> str: return hashlib.sha256(content).hexdigest()
def _name(headers: list[str]) -> str:
    joined = " ".join(headers).casefold()
    for term, name in (("balance", "Balance Sheet"), ("trial", "Trial Balance"), ("vendor", "Vendor Ledger"), ("invoice", "Invoice Details"), ("journal", "Journal Entries"), ("cash", "Cash Flow")):
        if term in joined: return name
    return "Financial Records"
def _table(doc_id: str, filename: str, rows: list[list[object]], table_name: str, page: int | None = None) -> FinancialTable:
    headers = [text(x) or f"Column {i+1}" for i, x in enumerate(rows[0])] if rows else []
    prov = Provenance(doc_id, filename, page, table_name)
    built=[]
    for i, row in enumerate(rows[1:], 1):
        cells=tuple(FinancialCell(text(v) or None, headers[j] if j < len(headers) else f"Column {j+1}", Provenance(doc_id, filename, page, table_name, i)) for j,v in enumerate(row))
        built.append(FinancialRow(cells, Provenance(doc_id, filename, page, table_name, i)))
    return FinancialTable(table_name, tuple(headers), tuple(built), prov)

@dataclass(slots=True)
class CsvParser:
    def parse(self, content: bytes, filename: str) -> FinancialDocument:
        doc_id=_id(content); sample=content[:4096].decode("utf-8-sig", errors="replace")
        dialect=csv.Sniffer().sniff(sample, delimiters=",;\t|")
        rows=list(csv.reader(io.StringIO(content.decode("utf-8-sig", errors="replace")), dialect))
        table=_table(doc_id, filename, rows, _name(rows[0] if rows else []))
        return FinancialDocument(doc_id, filename, "csv", {}, (FinancialSection(table.name,(table,),table.provenance),))

@dataclass(slots=True)
class ExcelParser:
    def parse(self, content: bytes, filename: str) -> FinancialDocument:
        doc_id=_id(content); book=pd.ExcelFile(io.BytesIO(content)); sections=[]
        for sheet in book.sheet_names:
            frame=pd.read_excel(book, sheet_name=sheet, header=None).fillna("")
            rows=frame.values.tolist()
            if rows:
                table=_table(doc_id,filename,rows,_name([str(x) for x in rows[0]]))
                sections.append(FinancialSection(sheet,(table,),table.provenance))
        return FinancialDocument(doc_id,filename,"excel",{},tuple(sections))

@dataclass(slots=True)
class DocxParser:
    def parse(self, content: bytes, filename: str) -> FinancialDocument:
        doc_id=_id(content); word=WordDocument(io.BytesIO(content)); tables=[]
        for table in word.tables:
            rows=[[cell.text for cell in row.cells] for row in table.rows]
            if rows: tables.append(_table(doc_id,filename,rows,_name(rows[0])))
        if not tables:
            rows=[["Text"], *[[p.text] for p in word.paragraphs if p.text.strip()]]; tables.append(_table(doc_id,filename,rows,"Document Details"))
        return FinancialDocument(doc_id,filename,"docx",{},tuple(FinancialSection(t.name,(t,),t.provenance) for t in tables))

@dataclass(slots=True)
class PdfParser:
    def parse(self, content: bytes, filename: str) -> FinancialDocument:
        doc_id=_id(content); tables=[]; warnings=[]
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page_no,page in enumerate(pdf.pages,1):
                extracted=page.extract_tables()
                if not extracted and not (page.extract_text() or "").strip(): warnings.append(f"Page {page_no}: image-only PDF; OCR unavailable")
                for rows in extracted:
                    clean=[[x or "" for x in row] for row in rows]
                    if clean: tables.append(_table(doc_id,filename,clean,_name(clean[0]),page_no))
        return FinancialDocument(doc_id,filename,"pdf",{},tuple(FinancialSection(t.name,(t,),t.provenance) for t in tables),tuple(warnings))

class ParserFactory:
    _parsers={"csv":CsvParser(),"xlsx":ExcelParser(),"xlsm":ExcelParser(),"xls":ExcelParser(),"docx":DocxParser(),"pdf":PdfParser()}
    @classmethod
    def parse(cls, content: bytes, filename: str) -> FinancialDocument:
        extension=filename.rsplit(".",1)[-1].casefold()
        if extension not in cls._parsers: raise ValueError(f"Unsupported file type: .{extension}")
        return cls._parsers[extension].parse(content,filename)
